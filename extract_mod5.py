import docx
import os

path = r"d:\agendamiento\correciones_Modulo5_Documentacion.docx"
text = ""
doc = docx.Document(path)
for para in doc.paragraphs:
    text += para.text + "\n"
for table in doc.tables:
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            # Avoid repeating text of merged cells
            row_text.append(cell.text.strip().replace("\n", " "))
        text += " | ".join(row_text) + "\n"

output_path = r"d:\agendamiento\extracted_modulo5.txt"
with open(output_path, "w", encoding="utf-8") as out:
    out.write(text)

print("Extraction complete!")
