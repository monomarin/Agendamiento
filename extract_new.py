import docx
import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

base = r"d:\agendamiento"
new_docx_files = [
    "correcionnueva_Modulo1_Definitivo (1).docx",
    "correcionnueva_Modulo2_Definitivo.docx",
    "correcionnueva_Modulo3_Definitivo.docx",
    "correcionnueva_Modulo4_Definitivo.docx",
]

def extract_docx(path):
    text = ""
    try:
        doc = docx.Document(path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                text += " | ".join(row_text) + "\n"
    except Exception as e:
        text = f"ERROR: {e}"
    return text

output_path = os.path.join(base, "extracted_new_docs.txt")
with open(output_path, "w", encoding="utf-8") as out:
    for f in new_docx_files:
        path = os.path.join(base, f)
        out.write(f"\n{'='*80}\n")
        out.write(f"FILE: {f}\n")
        out.write('='*80 + "\n")
        out.write(extract_docx(path))
        out.write("\n")
        print(f"Done: {f}")

print(f"\nAll extracted to: {output_path}")
